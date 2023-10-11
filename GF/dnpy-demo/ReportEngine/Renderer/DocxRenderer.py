import os

import docx
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


class DocxRenderer:

    def __init__(self, global_config):

        self.template_path = global_config['template_path']
        self.fig_path = global_config['fig_path']
        self.output_path = global_config['output_path']
        self.output_filename = global_config['output_filename']
        self.date = global_config['date']

        # 数据容器
        self.text_data_dict = {}  # 文本
        self.table_data_dict = {}  # 表格
        self.figure_title_dict = {}  # 图片标题

    # 解析标签
    def parse_label(self):
        pass

    # 渲染word文档，把文字、表格、图片填充上去
    def run(self):

        doc = docx.Document(self.template_path)

        doc.styles['Normal'].font.name = u'楷体'
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

        # -----4.1.修改报告日期
        doc.paragraphs[1].text = u'%d年%d月%d日' % (self.date.year, self.date.month, self.date.day)

        # -----4.2.替换正文文字
        for index, prg in enumerate(doc.paragraphs):
            text_list = prg.text.strip().split(' ')
            # 判断text的第一个词是否为<body>标签
            # 如果是正文文字段，第一个词为<body>标签，第二个词为文段的key
            if text_list[0] == '<body>':
                # 获取文段的key，再根据key从body_data_dict中获取对应文段
                text_key = text_list[1]
                text_value = self.text_data_dict.get(text_key, [])
                if text_value:
                    p = prg.clear()  # 清除已有内容
                    # 添加标题
                    run = p.add_run(text_value[0])
                    run.font.size = Pt(10.5)
                    run.font.bold = True
                    # 添加正文
                    run = p.add_run(text_value[1])
                    run.font.size = Pt(10.5)

        # -----4.2.替换表格或图片
        # 在模板中，图片也是用表格作为容器的（一个单元格放一张图片）。因此以下替换表格和图片的操作都和doc.tables有关
        for index, Table in enumerate(doc.tables):

            label = Table.cell(0, 0).text

            # -----4.2.1替换表格
            if label[0] == u'表':

                # 表3和表7的一些单独处理

                # 填充表3日期
                # if label[1] == '3':
                #     for i in range(1, 9, 2):
                #         Table.cell(0, 1).tables[0].cell(1, i).text = self.bonds_td
                #         Table.cell(0, 1).tables[0].cell(1, i).paragraphs[0].runs[0].font.size = 114300
                #         Table.cell(0, 1).tables[0].cell(1, i).paragraphs[
                #             0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # # 填充表7标题
                # elif label[1] == '7':
                #     title = self.figure_title_dict.get('table_7', u' ')
                #     run = Table.cell(0, 0).paragraphs[0].add_run(title)  # 修改文字
                #     run.font.size = Pt(10.5)  # 修改字体
                #     run.font.bold = True

                # 根据key获取表格数据
                table_key = 'tbl_' + label[1]
                new_data = self.table_data_dict.get(table_key, {})
                if len(new_data) != 0:
                    # 给表格填充数据，其中表7单独处理
                    if label[1] == '7':
                        tb = Table.cell(3, 0).tables[0]
                        for i, (_, row) in enumerate(new_data.iterrows()):
                            for j, d in enumerate(row):
                                tb.cell(i, j).text = d
                                tb.cell(i, j).paragraphs[0].runs[0].font.size = 114300
                                tb.cell(i, j).paragraphs[
                                    0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            tb.rows[i].height = Cm(0.46)

                    else:
                        for cell in Table.cell(0, 1).tables[0]._cells:
                            if len(cell.text.split('#')) == 2:
                                row, col = cell.text.split('#')
                                for r in cell.paragraphs[0].runs:
                                    r.text = ''
                                run = cell.paragraphs[0].add_run(new_data.loc[row.strip(), col.strip()])
                                run.font.size = 114300
                            else:
                                pass

            # -----4.2.2替换图片
            elif label[0] == u'图':
                # 每个单元格放一张图片
                for i in range(len(Table.rows)):
                    for j in range(len(Table.columns)):
                        cell = Table.cell(i, j)
                        cell_label = cell.text.split(u'-')

                        if cell_label[0] == u'<图>':

                            # -----4.2.2.1替换图标签
                            label = self.figure_title_dict.get('txt_title_fig%d' % (int(cell_label[1])), u' ')
                            run = Table.cell(i - 1, j).paragraphs[0].add_run(label)  # 修改文字
                            run.font.size = Pt(10.5)  # 修改字体
                            run.font.bold = True

                            # -----4.2.2.2替换图
                            path = os.path.join(self.fig_path, 'figure_%d.png' % (int(cell_label[1])))
                            if os.path.exists(path):
                                cell.paragraphs[0].text = u''
                                run = cell.paragraphs[0].runs[0].clear()

                                if int(cell_label[1]) in [5, 14, 15, 16, 19, 20, 23, 30, 33]:
                                    run.add_picture(path,
                                                    width=Inches(5.003937),
                                                    height=Inches(2.1299213))
                                else:
                                    run.add_picture(path, width=Inches(3.386),
                                                    height=Inches(2.03))
                            else:
                                pass

            else:
                pass

        doc.styles['Normal'].font.name = 'Times New Roman'

        doc.save(os.path.join(self.output_path, self.output_filename % (str(self.date).replace('-', ''))))
